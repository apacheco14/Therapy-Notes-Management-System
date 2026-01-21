#!/usr/bin/env python3
"""
Export Therapy Notes to Microsoft Word Documents

This script reads progress notes from a SQLite database and exports each note
as a separate, human-readable Word document (.docx).

Usage: python3 export_notes.py <database_path> [output_directory]
"""

import sqlite3
import os
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Color palette
class Colors:
    HEADER_TEXT = RGBColor(0x1A, 0x52, 0x76)
    LABEL_TEXT = RGBColor(0x5D, 0x6D, 0x7E)
    BORDER_COLOR = 'BDC3C7'
    HEADER_BG = 'E8F4F8'
    SECTION_BG = 'F8F9FA'
    CERTIFIED_GREEN = RGBColor(0x27, 0xAE, 0x60)
    NOT_CERTIFIED_RED = RGBColor(0xE7, 0x4C, 0x3C)


def format_date(date_str):
    """Format a date string for display with time."""
    if not date_str:
        return 'Not specified'
    try:
        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        return dt.strftime('%A, %B %d, %Y at %I:%M %p')
    except (ValueError, AttributeError):
        return str(date_str)


def format_simple_date(date_str):
    """Format a date string for display without time."""
    if not date_str:
        return 'Not specified'
    try:
        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        return dt.strftime('%B %d, %Y')
    except (ValueError, AttributeError):
        return str(date_str)


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color='BDC3C7'):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    
    tcPr.append(tcBorders)


def add_section_heading(doc, text):
    """Add a styled section heading."""
    para = doc.add_paragraph()
    para.space_before = Pt(15)
    para.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Arial'
    run.font.color.rgb = Colors.HEADER_TEXT
    return para


def add_label_value(doc, label, value, italics=False):
    """Add a label-value pair as a paragraph."""
    display_value = str(value) if value not in (None, '') else 'Not specified'
    
    para = doc.add_paragraph()
    para.space_before = Pt(3)
    para.space_after = Pt(3)
    
    label_run = para.add_run(f'{label}: ')
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.name = 'Arial'
    label_run.font.color.rgb = Colors.LABEL_TEXT
    
    value_run = para.add_run(display_value)
    value_run.font.size = Pt(11)
    value_run.font.name = 'Arial'
    value_run.italic = italics
    
    return para


def add_text_block(doc, label, text):
    """Add a multi-line text block with label."""
    # Label paragraph
    para = doc.add_paragraph()
    para.space_before = Pt(6)
    para.space_after = Pt(3)
    label_run = para.add_run(f'{label}:')
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.name = 'Arial'
    label_run.font.color.rgb = Colors.LABEL_TEXT
    
    # Content
    if text and text.strip():
        for line in text.split('\n'):
            content_para = doc.add_paragraph()
            content_para.paragraph_format.left_indent = Inches(0.25)
            content_para.space_before = Pt(3)
            content_para.space_after = Pt(3)
            content_run = content_para.add_run(line or ' ')
            content_run.font.size = Pt(11)
            content_run.font.name = 'Arial'
    else:
        content_para = doc.add_paragraph()
        content_para.paragraph_format.left_indent = Inches(0.25)
        content_run = content_para.add_run('Not specified')
        content_run.font.size = Pt(11)
        content_run.font.name = 'Arial'
        content_run.italic = True


def add_assessment_table(doc, items):
    """Create a table for mental status assessment items."""
    if not items:
        para = doc.add_paragraph()
        run = para.add_run('No items recorded')
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True
        return
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.0)
    
    # Header row
    header_row = table.rows[0]
    for idx, header_text in enumerate(['Assessment', 'Finding']):
        cell = header_row.cells[idx]
        cell.text = header_text
        set_cell_shading(cell, Colors.HEADER_BG)
        set_cell_borders(cell)
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    
    # Data rows
    for item in items:
        row = table.add_row()
        
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = item['label']
        set_cell_borders(label_cell)
        para = label_cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        # Value cell
        value_cell = row.cells[1]
        set_cell_borders(value_cell)
        
        # Clear default paragraph and add value
        value_cell.paragraphs[0].clear()
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(item['value'] or 'Not specified')
        value_run.font.size = Pt(11)
        value_run.font.name = 'Arial'
        
        # Add notes if present
        if item.get('notes'):
            notes_para = value_cell.add_paragraph()
            notes_label = notes_para.add_run('Notes: ')
            notes_label.font.size = Pt(10)
            notes_label.font.name = 'Arial'
            notes_label.italic = True
            notes_label.font.color.rgb = Colors.LABEL_TEXT
            
            notes_text = notes_para.add_run(item['notes'])
            notes_text.font.size = Pt(10)
            notes_text.font.name = 'Arial'
            notes_text.italic = True


def add_list_section(doc, title, items):
    """Add a section with a list of items in a table format."""
    add_section_heading(doc, title)
    
    if not items:
        para = doc.add_paragraph()
        para.space_after = Pt(6)
        run = para.add_run('None recorded')
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True
        return
    
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for idx, item in enumerate(items):
        row = table.add_row()
        cell = row.cells[0]
        set_cell_borders(cell)
        
        # Alternate row colors
        if idx % 2 == 1:
            set_cell_shading(cell, Colors.SECTION_BG)
        
        # Item name
        cell.paragraphs[0].clear()
        name_para = cell.paragraphs[0]
        name_run = name_para.add_run(item['name'])
        name_run.font.size = Pt(11)
        name_run.font.name = 'Arial'
        
        # Description if present
        if item.get('description'):
            desc_para = cell.add_paragraph()
            desc_run = desc_para.add_run(item['description'])
            desc_run.font.size = Pt(10)
            desc_run.font.name = 'Arial'
            desc_run.italic = True
            desc_run.font.color.rgb = Colors.LABEL_TEXT
        
        # Note if present
        if item.get('note'):
            note_para = cell.add_paragraph()
            note_label = note_para.add_run('Note: ')
            note_label.font.size = Pt(10)
            note_label.font.name = 'Arial'
            note_label.bold = True
            
            note_text = note_para.add_run(item['note'])
            note_text.font.size = Pt(10)
            note_text.font.name = 'Arial'


def generate_note_document(note, client, symptoms, referrals, collateral_contacts, assessment_options):
    """Generate a Word document for a single note."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Set margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== DOCUMENT HEADER =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    title_run = title.add_run('THERAPY PROGRESS NOTE')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = Colors.HEADER_TEXT
    
    # Certification status
    cert_para = doc.add_paragraph()
    cert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_para.space_after = Pt(15)
    
    if note['certified']:
        cert_text = f"Certified on {format_simple_date(note['certified'])}"
        cert_color = Colors.CERTIFIED_GREEN
    else:
        cert_text = 'Not yet certified'
        cert_color = Colors.NOT_CERTIFIED_RED
    
    cert_run = cert_para.add_run(cert_text)
    cert_run.font.size = Pt(10)
    cert_run.font.name = 'Arial'
    cert_run.italic = True
    cert_run.font.color.rgb = cert_color
    
    # ===== CLIENT INFORMATION =====
    add_section_heading(doc, 'Client Information')
    
    client_name = 'Unknown'
    if client:
        first = client['first_name'] or ''
        last = client['last_name'] or ''
        client_name = f'{first} {last}'.strip() or 'Unknown'
    
    add_label_value(doc, 'Client', client_name)
    add_label_value(doc, 'Client Code', client['client_code'] if client else None)
    if client and client.get('date_of_birth'):
        add_label_value(doc, 'Date of Birth', format_simple_date(client['date_of_birth']))
    
    # ===== SESSION INFORMATION =====
    add_section_heading(doc, 'Session Information')
    
    add_label_value(doc, 'Appointment Date/Time', format_date(note['appt_date_time']))
    add_label_value(doc, 'Session Type', 'Virtual/Telehealth' if note['virtual_appt'] else 'In-Person')
    add_label_value(doc, 'Session Number', note['session_number'])
    add_label_value(doc, 'Session Length', note['session_length'])
    add_label_value(doc, 'Diagnosis', note['diagnosis'])
    
    if note['appt_note']:
        add_text_block(doc, 'Appointment Notes', note['appt_note'])
    
    # ===== NARRATIVE =====
    if note['narrative']:
        add_section_heading(doc, 'Session Narrative')
        for para_text in note['narrative'].split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph()
                para.space_after = Pt(6)
                run = para.add_run(para_text.strip())
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    # ===== SYMPTOMS =====
    symptom_items = [{'name': s['name'], 'description': s['description']} for s in symptoms]
    add_list_section(doc, 'Presenting Symptoms', symptom_items)
    
    # ===== MENTAL STATUS ASSESSMENT =====
    add_section_heading(doc, 'Mental Status Assessment')
    
    def get_option_name(option_id):
        if not option_id:
            return None
        for opt in assessment_options:
            if opt['id'] == option_id:
                return opt['name']
        return None
    
    assessment_items = [
        {'label': 'Appearance', 'value': get_option_name(note['appearance']), 'notes': note['appearance_notes']},
        {'label': 'Speech', 'value': get_option_name(note['speech']), 'notes': note['speech_notes']},
        {'label': 'Affect', 'value': get_option_name(note['affect']), 'notes': note['affect_notes']},
        {'label': 'Eye Contact', 'value': get_option_name(note['eye_contact']), 'notes': note['eye_contact_notes']}
    ]
    
    add_assessment_table(doc, assessment_items)
    
    # ===== COLLATERAL CONTACTS =====
    if collateral_contacts:
        contact_items = [
            {'name': c['name'] or 'Contact', 'description': c['description'], 'note': c['collateral_contact_note']}
            for c in collateral_contacts
        ]
        add_list_section(doc, 'Collateral Contacts', contact_items)
    
    # ===== REFERRALS =====
    if referrals:
        referral_items = [
            {'name': r['name'], 'description': r['description'], 'note': r['referral_note']}
            for r in referrals
        ]
        add_list_section(doc, 'Referrals Made', referral_items)
    
    # ===== NEXT APPOINTMENT =====
    add_section_heading(doc, 'Follow-Up')
    
    next_appt_option = get_option_name(note['next_appt'])
    add_label_value(doc, 'Next Appointment', next_appt_option)
    if note['next_appt_notes']:
        add_text_block(doc, 'Follow-Up Notes', note['next_appt_notes'])
    
    # ===== FOOTER =====
    doc.add_paragraph()  # Spacer
    
    divider = doc.add_paragraph()
    divider.space_before = Pt(20)
    divider_run = divider.add_run('─' * 60)
    divider_run.font.size = Pt(10)
    divider_run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
    
    footer = doc.add_paragraph()
    footer.space_before = Pt(6)
    footer_text = f"Note ID: {note['note_id']}  |  Created: {format_simple_date(note['insert_date'])}  |  Last Updated: {format_simple_date(note['update_date'])}"
    footer_run = footer.add_run(footer_text)
    footer_run.font.size = Pt(9)
    footer_run.font.name = 'Arial'
    footer_run.font.color.rgb = Colors.LABEL_TEXT
    
    return doc


def dict_factory(cursor, row):
    """Convert sqlite row to dictionary."""
    return {col[0]: row[idx] for idx, col in enumerate(cursor.description)}


def export_notes(db_path, output_dir):
    """Main export function."""
    # Validate database path
    if not os.path.exists(db_path):
        print(f'Error: Database file not found: {db_path}')
        sys.exit(1)
    
    # Create output directory if needed
    os.makedirs(output_dir, exist_ok=True)
    
    # Connect to database
    conn = sqlite3.connect(db_path)
    conn.row_factory = dict_factory
    cursor = conn.cursor()
    
    try:
        # Get all notes
        cursor.execute('SELECT * FROM notes ORDER BY appt_date_time DESC')
        notes = cursor.fetchall()
        
        if not notes:
            print('No notes found in database.')
            return
        
        print(f'Found {len(notes)} note(s) to export.')
        
        # Pre-fetch all assessment options
        cursor.execute('SELECT id, type, name, description FROM assessment_options')
        assessment_options = cursor.fetchall()
        
        # Process each note
        for note in notes:
            print(f"Processing note ID {note['note_id']}...")
            
            # Get client information
            cursor.execute('SELECT * FROM clients WHERE client_id = ?', (note['client_id'],))
            client = cursor.fetchone()
            
            # Get symptoms for this note
            cursor.execute('''
                SELECT ao.name, ao.description
                FROM symptoms s
                JOIN assessment_options ao ON s.symptom_id = ao.id
                WHERE s.note_id = ?
            ''', (note['note_id'],))
            symptoms = cursor.fetchall()
            
            # Get referrals for this note
            cursor.execute('''
                SELECT ao.name, ao.description, r.referral_note
                FROM referrals r
                JOIN assessment_options ao ON r.referral_id = ao.id
                WHERE r.note_id = ?
            ''', (note['note_id'],))
            referrals = cursor.fetchall()
            
            # Get collateral contacts for this note
            cursor.execute('''
                SELECT ao.name, ao.description, cc.collateral_contact_note
                FROM collateral_contacts cc
                JOIN assessment_options ao ON cc.collateral_contact_type_id = ao.id
                WHERE cc.note_id = ?
            ''', (note['note_id'],))
            collateral_contacts = cursor.fetchall()
            
            # Generate document
            doc = generate_note_document(
                note, client, symptoms, referrals, collateral_contacts, assessment_options
            )
            
            # Create filename
            client_code = client['client_code'] if client else 'unknown'
            date_str = 'no-date'
            if note['appt_date_time']:
                try:
                    dt = datetime.fromisoformat(note['appt_date_time'].replace('Z', '+00:00'))
                    date_str = dt.strftime('%Y-%m-%d')
                except ValueError:
                    pass
            
            filename = f"note_{note['note_id']}_{client_code}_{date_str}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # Save document
            doc.save(filepath)
            print(f'  Exported: {filename}')
        
        print(f'\nExport complete! {len(notes)} document(s) saved to: {output_dir}')
        
    finally:
        conn.close()


def main():
    """CLI entry point."""
    args = sys.argv[1:]
    
    if len(args) < 1:
        print('''
Therapy Notes Export Tool
=========================

Usage: python3 export_notes.py <database_path> [output_directory]

Arguments:
  database_path     Path to the SQLite database file
  output_directory  Directory for exported .docx files (default: ./exported_notes)

Example:
  python3 export_notes.py therapy.db ./exports
''')
        sys.exit(0)
    
    db_path = args[0]
    output_dir = args[1] if len(args) > 1 else './exported_notes'
    
    export_notes(db_path, output_dir)


if __name__ == '__main__':
    main()